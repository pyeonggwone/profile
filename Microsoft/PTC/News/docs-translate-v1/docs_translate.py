# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "litellm>=1.50",
#   "openai>=1.40",
#   "pydantic-settings>=2.5",
#   "python-docx>=1.1.2",
#   "typer>=0.12",
#   "rich>=13.8",
# ]
# ///
from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import sqlite3
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import typer
from pydantic_settings import BaseSettings, SettingsConfigDict
from rich.console import Console

if sys.version_info < (3, 11):
    raise SystemExit("Python 3.11+ 이 필요합니다.")

try:
    import litellm
except Exception:  # pragma: no cover - handled at runtime
    litellm = None

try:
    from docx import Document
except Exception:  # pragma: no cover - handled at runtime
    Document = None

app = typer.Typer(no_args_is_help=True, add_completion=False)
tm_app = typer.Typer(no_args_is_help=True, add_completion=False)
app.add_typer(tm_app, name="tm")
console = Console()

SUPPORTED_LANGS = {"en", "kr", "ko", "ch", "zh", "jp", "ja"}
SUPPORTED_EXTENSIONS = {".md", ".docx"}
TARGET_SUFFIX = {"kr": "KR", "ko": "KR", "ch": "CH", "zh": "CH", "jp": "JP", "ja": "JP", "en": "EN"}
LANG_NAME = {
    "en": "English",
    "kr": "Korean",
    "ko": "Korean",
    "ch": "Chinese",
    "zh": "Chinese",
    "jp": "Japanese",
    "ja": "Japanese",
}
BATCH_SIZE = 8
PLACEHOLDER_RE = re.compile(r"__DOCSTR_[A-Z]+_\d{4}__")
BLOCK_MARKER_RE = re.compile(r"^===\s*(\d+)\s*===\s*$", re.MULTILINE)
FENCE_RE = re.compile(r"^\s*(```|~~~)")
HEADING_RE = re.compile(r"^(\s{0,3}#{1,6}\s+)(.*?)(\s+#+\s*)?$")
LIST_RE = re.compile(r"^(\s*(?:(?:[-+*])|(?:\d+[.)]))\s+(?:\[[ xX]\]\s+)?)(.+?)\s*$")
BLOCKQUOTE_RE = re.compile(r"^(\s*>\s?)(.+?)\s*$")
ADMONITION_RE = re.compile(r"^(\s*>?\s*!!!\s+\w+\s*)(.*)$")
HTML_COMMENT_RE = re.compile(r"<!--.*?-->")
HTML_TAG_RE = re.compile(r"</?[^>\n]+?>")
INLINE_CODE_RE = re.compile(r"`[^`\n]+`")
MATH_RE = re.compile(r"\$[^$\n]+\$")
LIQUID_RE = re.compile(r"({[{%#].*?[}%]})")
ANGLE_VAR_RE = re.compile(r"<%.*?%>")
BRACE_VAR_RE = re.compile(r"\$\{[^}\n]+\}")
URL_RE = re.compile(r"(?<!\()\bhttps?://[^\s)\]}>\"]+")
MD_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
MD_LINK_RE = re.compile(r"(?<!!)\[([^\]]+)\]\(([^)]+)\)")
REF_LINK_RE = re.compile(r"(?<!!)\[([^\]]+)\]\[([^\]]*)\]")
AUTOLINK_RE = re.compile(r"<https?://[^>]+>")


class Settings(BaseSettings):
    model_config = SettingsConfigDict(env_file=".env", extra="ignore")

    openai_api_key: str | None = None
    openai_model: str = "gpt-4o-mini"
    azure_openai_api_key: str | None = None
    azure_openai_endpoint: str | None = None
    azure_openai_api_version: str = "2024-08-01-preview"
    azure_openai_deployment: str | None = None
    source_lang: str = "en"
    target_lang: str = "kr"
    work_dir: str = "work"
    tm_db_path: str = "work/tm.sqlite"
    glossary_path: str = "glossary.csv"

    @property
    def llm_model(self) -> str:
        if self.azure_openai_deployment:
            return f"azure/{self.azure_openai_deployment}"
        return self.openai_model


def settings() -> Settings:
    cfg = Settings()
    if cfg.openai_api_key:
        os.environ["OPENAI_API_KEY"] = cfg.openai_api_key
    if cfg.azure_openai_api_key:
        os.environ["AZURE_API_KEY"] = cfg.azure_openai_api_key
        os.environ["AZURE_OPENAI_API_KEY"] = cfg.azure_openai_api_key
    if cfg.azure_openai_endpoint:
        os.environ["AZURE_API_BASE"] = cfg.azure_openai_endpoint
        os.environ["AZURE_OPENAI_ENDPOINT"] = cfg.azure_openai_endpoint
    if cfg.azure_openai_api_version:
        os.environ["AZURE_API_VERSION"] = cfg.azure_openai_api_version
        os.environ["AZURE_OPENAI_API_VERSION"] = cfg.azure_openai_api_version
    return cfg


@dataclass
class ProtectedToken:
    id: str
    value: str
    kind: str


class TokenMasker:
    def __init__(self, protected_terms: list[str] | None = None) -> None:
        self.counter = 0
        self.protected_terms = protected_terms or []

    def _token(self, kind: str, value: str, tokens: list[dict[str, str]]) -> str:
        self.counter += 1
        token = f"__DOCSTR_{kind.upper()}_{self.counter:04d}__"
        tokens.append({"id": token, "value": value, "kind": kind})
        return token

    def _sub(self, pattern: re.Pattern[str], text: str, tokens: list[dict[str, str]], kind: str) -> str:
        return pattern.sub(lambda match: self._token(kind, match.group(0), tokens), text)

    def mask(self, text: str) -> tuple[str, list[dict[str, str]]]:
        tokens: list[dict[str, str]] = []
        masked = text

        def image_repl(match: re.Match[str]) -> str:
            alt = match.group(1)
            target = match.group(2)
            return f"![{alt}]({self._token('imagepath', target, tokens)})"

        def link_repl(match: re.Match[str]) -> str:
            label = match.group(1)
            target = match.group(2)
            return f"[{label}]({self._token('linkurl', target, tokens)})"

        def ref_link_repl(match: re.Match[str]) -> str:
            label = match.group(1)
            ref = match.group(2)
            return f"[{label}][{self._token('ref', ref, tokens)}]"

        masked = MD_IMAGE_RE.sub(image_repl, masked)
        masked = MD_LINK_RE.sub(link_repl, masked)
        masked = REF_LINK_RE.sub(ref_link_repl, masked)
        masked = self._sub(AUTOLINK_RE, masked, tokens, "autolink")
        masked = self._sub(URL_RE, masked, tokens, "url")
        masked = self._sub(INLINE_CODE_RE, masked, tokens, "code")
        masked = self._sub(MATH_RE, masked, tokens, "math")
        masked = self._sub(LIQUID_RE, masked, tokens, "template")
        masked = self._sub(ANGLE_VAR_RE, masked, tokens, "template")
        masked = self._sub(BRACE_VAR_RE, masked, tokens, "variable")
        masked = self._sub(HTML_COMMENT_RE, masked, tokens, "html_comment")
        masked = self._sub(HTML_TAG_RE, masked, tokens, "html_tag")

        for term in sorted(self.protected_terms, key=len, reverse=True):
            if not term:
                continue
            pattern = re.compile(rf"(?<!\w){re.escape(term)}(?!\w)")
            masked = pattern.sub(lambda match: self._token("term", match.group(0), tokens), masked)
        return masked, tokens


def restore_tokens(text: str, tokens: list[dict[str, str]]) -> str:
    restored = text
    for token in tokens:
        restored = restored.replace(token["id"], token["value"])
    return restored


def normalize_lang(value: str) -> str:
    lowered = value.lower()
    if lowered not in SUPPORTED_LANGS:
        raise typer.BadParameter(f"지원 언어가 아닙니다: {value}. en/kr/ch/jp 중 하나를 사용하세요.")
    if lowered == "ko":
        return "kr"
    if lowered == "zh":
        return "ch"
    if lowered == "ja":
        return "jp"
    return lowered


def load_glossary(path: Path) -> tuple[list[dict[str, str]], list[str]]:
    if not path.exists():
        return [], []
    rows: list[dict[str, str]] = []
    protected: list[str] = []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            term = (row.get("term") or "").strip()
            translation = (row.get("translation") or "").strip()
            is_protected = (row.get("protected") or "").strip().lower() == "true"
            if not term:
                continue
            rows.append({"term": term, "translation": translation, "protected": str(is_protected).lower()})
            if is_protected:
                protected.append(term)
    return rows, protected


def work_key(path: Path) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", str(path.with_suffix("")))[:180]


def default_work_file(input_path: Path, filename: str, cfg: Settings) -> Path:
    return Path(cfg.work_dir) / work_key(input_path) / filename


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def is_reference_definition(line: str) -> bool:
    return bool(re.match(r"^\s{0,3}\[[^\]]+\]:\s+", line))


def split_table_row(line: str) -> list[str]:
    return line.rstrip("\n").split("|")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if "|" not in stripped:
        return False
    parts = [part.strip() for part in stripped.strip("|").split("|")]
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", part or "") for part in parts)


def should_skip_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return True
    if is_table_separator(line):
        return True
    if is_reference_definition(line):
        return True
    if stripped.startswith(("<!--", "[//]:", "{:")):
        return True
    return False


def segment_for_line(
    line: str,
    line_index: int,
    masker: TokenMasker,
    current_heading: str | None,
) -> list[dict[str, Any]]:
    if should_skip_line(line):
        return []
    raw = line.rstrip("\n")
    segments: list[dict[str, Any]] = []

    heading = HEADING_RE.match(raw)
    if heading:
        text = heading.group(2).strip()
        if text:
            masked, tokens = masker.mask(text)
            segments.append({
                "id": len(segments) + 1,
                "path": ["line", line_index, "heading"],
                "type": "heading",
                "text": masked,
                "raw_text": text,
                "tokens": tokens,
                "context": {"heading": current_heading, "level": heading.group(1).count("#")},
            })
        return segments

    admonition = ADMONITION_RE.match(raw)
    if admonition and admonition.group(2).strip():
        text = admonition.group(2).strip().strip('"')
        masked, tokens = masker.mask(text)
        segments.append({
            "id": 1,
            "path": ["line", line_index, "admonition_title"],
            "type": "admonition_title",
            "text": masked,
            "raw_text": text,
            "tokens": tokens,
            "context": {"heading": current_heading},
        })
        return segments

    if "|" in raw and not is_table_separator(raw):
        cells = split_table_row(raw)
        for cell_index, cell in enumerate(cells):
            text = cell.strip()
            if not text or re.fullmatch(r":?-{3,}:?", text):
                continue
            masked, tokens = masker.mask(text)
            segments.append({
                "id": len(segments) + 1,
                "path": ["line", line_index, "table_cell", cell_index],
                "type": "table_cell",
                "text": masked,
                "raw_text": text,
                "tokens": tokens,
                "context": {"heading": current_heading},
            })
        return segments

    blockquote = BLOCKQUOTE_RE.match(raw)
    if blockquote:
        text = blockquote.group(2).strip()
        if text:
            masked, tokens = masker.mask(text)
            segments.append({
                "id": 1,
                "path": ["line", line_index, "blockquote"],
                "type": "blockquote",
                "text": masked,
                "raw_text": text,
                "tokens": tokens,
                "context": {"heading": current_heading},
            })
        return segments

    list_item = LIST_RE.match(raw)
    if list_item:
        text = list_item.group(2).strip()
        if text:
            masked, tokens = masker.mask(text)
            segments.append({
                "id": 1,
                "path": ["line", line_index, "list_item"],
                "type": "list_item",
                "text": masked,
                "raw_text": text,
                "tokens": tokens,
                "context": {"heading": current_heading},
            })
        return segments

    masked, tokens = masker.mask(raw.strip())
    if masked.strip():
        segments.append({
            "id": 1,
            "path": ["line", line_index, "paragraph"],
            "type": "paragraph",
            "text": masked,
            "raw_text": raw.strip(),
            "tokens": tokens,
            "context": {"heading": current_heading},
        })
    return segments


def extract_segments(input_path: Path, cfg: Settings) -> dict[str, Any]:
    if input_path.suffix.lower() == ".docx":
        return extract_docx_segments(input_path, cfg)
    return extract_markdown_segments(input_path, cfg)


def extract_markdown_segments(input_path: Path, cfg: Settings) -> dict[str, Any]:
    glossary, protected_terms = load_glossary(Path(cfg.glossary_path))
    masker = TokenMasker(protected_terms)
    text = read_text(input_path)
    lines = text.splitlines(keepends=True)
    segments: list[dict[str, Any]] = []
    in_fence = False
    in_frontmatter = False
    current_heading: str | None = None

    for line_index, line in enumerate(lines):
        stripped = line.strip()
        if line_index == 0 and stripped == "---":
            in_frontmatter = True
            continue
        if in_frontmatter:
            if stripped == "---":
                in_frontmatter = False
            continue
        if FENCE_RE.match(line):
            in_fence = not in_fence
            continue
        if in_fence:
            continue

        line_segments = segment_for_line(line, line_index, masker, current_heading)
        for item in line_segments:
            item["id"] = len(segments) + 1
            item["file"] = str(input_path)
            segments.append(item)

        heading = HEADING_RE.match(line.rstrip("\n"))
        if heading and heading.group(2).strip():
            current_heading = heading.group(2).strip()

    return {
        "version": 1,
        "format": "markdown",
        "source_file": str(input_path),
        "source_lang": normalize_lang(cfg.source_lang),
        "target_lang": normalize_lang(cfg.target_lang),
        "glossary": glossary,
        "segments": segments,
    }


def extract_docx_segments(input_path: Path, cfg: Settings) -> dict[str, Any]:
    if Document is None:
        raise RuntimeError(".docx 처리를 위해 python-docx 가 필요합니다.")
    glossary, protected_terms = load_glossary(Path(cfg.glossary_path))
    masker = TokenMasker(protected_terms)
    document = Document(str(input_path))
    segments: list[dict[str, Any]] = []

    def add_segment(path: list[Any], segment_type: str, text: str) -> None:
        clean = text.strip()
        if not clean:
            return
        masked, tokens = masker.mask(clean)
        segments.append({
            "id": len(segments) + 1,
            "file": str(input_path),
            "path": path,
            "type": segment_type,
            "text": masked,
            "raw_text": clean,
            "tokens": tokens,
            "context": {},
        })

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        add_segment(["paragraph", paragraph_index], "docx_paragraph", paragraph.text)

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    add_segment(["table", table_index, row_index, cell_index, paragraph_index], "docx_table_cell", paragraph.text)

    return {
        "version": 1,
        "format": "docx",
        "source_file": str(input_path),
        "source_lang": normalize_lang(cfg.source_lang),
        "target_lang": normalize_lang(cfg.target_lang),
        "glossary": glossary,
        "segments": segments,
    }


def init_tm(db_path: Path) -> None:
    db_path.parent.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS tm (
                src_hash TEXT PRIMARY KEY,
                src TEXT NOT NULL,
                tgt TEXT NOT NULL,
                model TEXT,
                source_lang TEXT,
                target_lang TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """
        )


def hash_src(src: str, source_lang: str, target_lang: str) -> str:
    payload = f"{source_lang}\n{target_lang}\n{src}"
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def tm_get(db_path: Path, src: str, source_lang: str, target_lang: str) -> str | None:
    init_tm(db_path)
    src_hash = hash_src(src, source_lang, target_lang)
    with sqlite3.connect(db_path) as conn:
        row = conn.execute("SELECT tgt FROM tm WHERE src_hash = ?", (src_hash,)).fetchone()
    return row[0] if row else None


def tm_put(db_path: Path, src: str, tgt: str, model: str, source_lang: str, target_lang: str) -> None:
    init_tm(db_path)
    src_hash = hash_src(src, source_lang, target_lang)
    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            INSERT OR REPLACE INTO tm(src_hash, src, tgt, model, source_lang, target_lang)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (src_hash, src, tgt, model, source_lang, target_lang),
        )


def glossary_prompt(glossary: list[dict[str, str]]) -> str:
    if not glossary:
        return "(none)"
    lines = []
    for row in glossary:
        term = row["term"]
        translation = row.get("translation") or term
        protected = row.get("protected") == "true"
        suffix = " (protected, keep exactly)" if protected else ""
        lines.append(f"- {term} => {translation}{suffix}")
    return "\n".join(lines)


def parse_output_block(text: str, expected: int) -> list[str] | None:
    matches = list(BLOCK_MARKER_RE.finditer(text))
    if len(matches) != expected:
        return None
    results: list[str] = []
    for idx, match in enumerate(matches):
        number = int(match.group(1))
        if number != idx + 1:
            return None
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        results.append(text[start:end].strip())
    return results


def call_llm_block(items: list[dict[str, Any]], payload: dict[str, Any], cfg: Settings) -> list[str]:
    if litellm is None:
        raise RuntimeError("litellm 을 import 할 수 없습니다. requirements.txt 를 설치하세요.")
    source_lang = normalize_lang(payload.get("source_lang", cfg.source_lang))
    target_lang = normalize_lang(payload.get("target_lang", cfg.target_lang))
    source_name = LANG_NAME[source_lang]
    target_name = LANG_NAME[target_lang]
    glossary = glossary_prompt(payload.get("glossary", []))
    system = f"""You are a professional technical documentation translator.
Translate from {source_name} to {target_name}.
Return only numbered blocks in the exact same format.
Do NOT add JSON, bullets, commentary, or explanations.
Preserve every placeholder matching __DOCSTR_*_0000__ exactly.
Preserve Markdown syntax around placeholders exactly.
Translate link text, headings, paragraphs, list items, table cells, and blockquote text naturally.
Do not translate code, URLs, variables, HTML tags, or protected glossary terms represented as placeholders.
Glossary:
{glossary}
"""
    user_lines: list[str] = []
    for index, item in enumerate(items, start=1):
        item_type = item.get("type", "paragraph")
        context = item.get("context") or {}
        heading = context.get("heading")
        user_lines.append(f"=== {index} ===")
        user_lines.append(f"[type: {item_type}; heading: {heading or 'none'}]")
        user_lines.append(item["text"])
    response = litellm.completion(
        model=cfg.llm_model,
        temperature=0,
        max_tokens=4096,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": "\n".join(user_lines)}],
    )
    content = response.choices[0].message.content or ""
    parsed = parse_output_block(content, len(items))
    if parsed is None:
        raise ValueError("LLM 응답 번호 블록 파싱 실패")
    cleaned: list[str] = []
    for translated in parsed:
        cleaned.append(re.sub(r"^\[type:.*?\]\s*", "", translated, flags=re.DOTALL).strip())
    return cleaned


def translate_payload(payload: dict[str, Any], cfg: Settings) -> dict[str, Any]:
    source_lang = normalize_lang(payload.get("source_lang", cfg.source_lang))
    target_lang = normalize_lang(payload.get("target_lang", cfg.target_lang))
    db_path = Path(cfg.tm_db_path)
    translated_segments: list[dict[str, Any]] = []
    misses: list[dict[str, Any]] = []

    for segment in payload["segments"]:
        cached = tm_get(db_path, segment["text"], source_lang, target_lang)
        item = dict(segment)
        if cached is not None:
            item["translated_text"] = cached
            item["tm"] = "hit"
            translated_segments.append(item)
        else:
            item["tm"] = "miss"
            translated_segments.append(item)
            misses.append(item)

    total = len(misses)
    for start in range(0, total, BATCH_SIZE):
        batch = misses[start : start + BATCH_SIZE]
        try:
            outputs = call_llm_block(batch, payload, cfg)
        except Exception as exc:
            console.print(f"[yellow]LLM 배치 실패: {exc}. 원문 fallback 처리[/yellow]")
            outputs = [item["text"] for item in batch]
        for item, translated in zip(batch, outputs, strict=True):
            source_placeholders = sorted(PLACEHOLDER_RE.findall(item["text"]))
            target_placeholders = sorted(PLACEHOLDER_RE.findall(translated))
            if source_placeholders != target_placeholders:
                console.print(f"[yellow]placeholder 불일치: segment {item['id']} 원문 fallback[/yellow]")
                translated = item["text"]
            tm_put(db_path, item["text"], translated, cfg.llm_model, source_lang, target_lang)
            for saved in translated_segments:
                if saved["id"] == item["id"]:
                    saved["translated_text"] = translated
                    break
        console.print(f"TRANSLATE {min(start + BATCH_SIZE, total)}/{total}")

    return {
        "version": payload.get("version", 1),
        "format": payload.get("format", "markdown"),
        "source_file": payload.get("source_file"),
        "source_lang": source_lang,
        "target_lang": target_lang,
        "model": cfg.llm_model,
        "segments": translated_segments,
    }


def replace_line_segment(line: str, segment: dict[str, Any], translated_text: str) -> str:
    raw = line.rstrip("\n")
    newline = "\n" if line.endswith("\n") else ""
    path = segment["path"]
    kind = path[2]
    restored = restore_tokens(translated_text, segment.get("tokens", []))

    if kind == "heading":
        match = HEADING_RE.match(raw)
        if match:
            return f"{match.group(1)}{restored}{match.group(3) or ''}{newline}"
    if kind == "admonition_title":
        match = ADMONITION_RE.match(raw)
        if match:
            return f"{match.group(1)}\"{restored}\"{newline}"
    if kind == "table_cell":
        cell_index = path[3]
        cells = split_table_row(raw)
        if 0 <= cell_index < len(cells):
            original = cells[cell_index]
            leading = original[: len(original) - len(original.lstrip())]
            trailing = original[len(original.rstrip()) :]
            cells[cell_index] = f"{leading}{restored}{trailing}"
            return "|".join(cells) + newline
    if kind == "blockquote":
        match = BLOCKQUOTE_RE.match(raw)
        if match:
            return f"{match.group(1)}{restored}{newline}"
    if kind == "list_item":
        match = LIST_RE.match(raw)
        if match:
            return f"{match.group(1)}{restored}{newline}"
    if kind == "paragraph":
        leading = raw[: len(raw) - len(raw.lstrip())]
        trailing = raw[len(raw.rstrip()) :]
        return f"{leading}{restored}{trailing}{newline}"
    return line


def apply_translations(input_path: Path, translated_path: Path, output_path: Path) -> dict[str, Any]:
    payload = read_json(translated_path)
    if payload.get("format") == "docx" or input_path.suffix.lower() == ".docx":
        return apply_docx_translations(input_path, payload, output_path)
    return apply_markdown_translations(input_path, payload, output_path)


def apply_markdown_translations(input_path: Path, payload: dict[str, Any], output_path: Path) -> dict[str, Any]:
    lines = read_text(input_path).splitlines(keepends=True)
    by_line: dict[int, list[dict[str, Any]]] = {}
    for segment in payload["segments"]:
        translated_text = segment.get("translated_text", segment.get("text", ""))
        segment = dict(segment)
        segment["translated_text"] = translated_text
        by_line.setdefault(segment["path"][1], []).append(segment)

    for line_index, items in by_line.items():
        if not (0 <= line_index < len(lines)):
            continue
        if len(items) == 1:
            lines[line_index] = replace_line_segment(lines[line_index], items[0], items[0]["translated_text"])
        else:
            table_items = [item for item in items if item["path"][2] == "table_cell"]
            if table_items:
                current = lines[line_index]
                for item in table_items:
                    current = replace_line_segment(current, item, item["translated_text"])
                lines[line_index] = current

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("".join(lines), encoding="utf-8")
    return verify_output(output_path)


def set_paragraph_text(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def apply_docx_translations(input_path: Path, payload: dict[str, Any], output_path: Path) -> dict[str, Any]:
    if Document is None:
        raise RuntimeError(".docx 처리를 위해 python-docx 가 필요합니다.")
    document = Document(str(input_path))
    for segment in payload["segments"]:
        translated_text = segment.get("translated_text", segment.get("text", ""))
        restored = restore_tokens(translated_text, segment.get("tokens", []))
        path = segment["path"]
        if path[0] == "paragraph":
            paragraph_index = path[1]
            if 0 <= paragraph_index < len(document.paragraphs):
                set_paragraph_text(document.paragraphs[paragraph_index], restored)
        elif path[0] == "table":
            table_index, row_index, cell_index, paragraph_index = path[1:]
            if table_index < len(document.tables):
                table = document.tables[table_index]
                if row_index < len(table.rows) and cell_index < len(table.rows[row_index].cells):
                    cell = table.rows[row_index].cells[cell_index]
                    if paragraph_index < len(cell.paragraphs):
                        set_paragraph_text(cell.paragraphs[paragraph_index], restored)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return verify_docx_output(output_path)


def verify_output(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8")
    issues: list[str] = []
    if PLACEHOLDER_RE.search(text):
        issues.append("unrestored_placeholder")
    fence_count = sum(1 for line in text.splitlines() if FENCE_RE.match(line))
    if fence_count % 2 != 0:
        issues.append("unbalanced_code_fence")
    return {"ok": not issues, "issues": issues, "path": str(path)}


def verify_docx_output(path: Path) -> dict[str, Any]:
    if Document is None:
        return {"ok": False, "issues": ["python_docx_unavailable"], "path": str(path)}
    document = Document(str(path))
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.extend(paragraph.text for paragraph in cell.paragraphs)
    joined = "\n".join(text_parts)
    issues: list[str] = []
    if PLACEHOLDER_RE.search(joined):
        issues.append("unrestored_placeholder")
    return {"ok": not issues, "issues": issues, "path": str(path)}


def default_output_path(input_path: Path, target_lang: str, base_input: Path | None = None) -> Path:
    suffix = TARGET_SUFFIX.get(target_lang, target_lang.upper())
    if base_input and base_input.is_dir():
        try:
            rel = input_path.relative_to(base_input)
            return Path("output") / rel.parent / f"{input_path.stem}_{suffix}{input_path.suffix}"
        except ValueError:
            pass
    if input_path.parent.name.lower() == "input":
        return input_path.parent.parent / "output" / f"{input_path.stem}_{suffix}{input_path.suffix}"
    return input_path.with_name(f"{input_path.stem}_{suffix}{input_path.suffix}")


def discover_inputs(path: Path) -> list[Path]:
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED_EXTENSIONS else []
    files = []
    for item in path.rglob("*"):
        if not item.is_file() or item.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        parts = {part.lower() for part in item.parts}
        if "done" in parts or "output" in parts or "work" in parts:
            continue
        files.append(item)
    return sorted(files)


def move_done(path: Path) -> Path:
    done_dir = path.parent / "done"
    done_dir.mkdir(parents=True, exist_ok=True)
    target = done_dir / path.name
    if target.exists():
        stamp = datetime.now().strftime("%Y%m%d%H%M%S")
        target = done_dir / f"{path.stem}_{stamp}{path.suffix}"
    shutil.move(str(path), str(target))
    return target


@app.command()
def extract(
    input_path: Path = typer.Argument(..., exists=True, readable=True),
    out: Path | None = typer.Option(None, "--out"),
    in_lang: str | None = typer.Option(None, "-in_lang", "--in-lang"),
    out_lang: str | None = typer.Option(None, "-out_lang", "--out-lang"),
) -> None:
    cfg = settings()
    if in_lang:
        cfg.source_lang = normalize_lang(in_lang)
    if out_lang:
        cfg.target_lang = normalize_lang(out_lang)
    payload = extract_segments(input_path, cfg)
    output = out or default_work_file(input_path, "segments.json", cfg)
    write_json(output, payload)
    console.print(f"EXTRACT {len(payload['segments'])} segments -> {output}")


@app.command()
def translate(
    segments_path: Path = typer.Argument(..., exists=True, readable=True),
    out: Path | None = typer.Option(None, "--out"),
    in_lang: str | None = typer.Option(None, "-in_lang", "--in-lang"),
    out_lang: str | None = typer.Option(None, "-out_lang", "--out-lang"),
) -> None:
    cfg = settings()
    payload = read_json(segments_path)
    if in_lang:
        payload["source_lang"] = normalize_lang(in_lang)
    if out_lang:
        payload["target_lang"] = normalize_lang(out_lang)
    translated = translate_payload(payload, cfg)
    output = out or segments_path.with_name("translated.json")
    write_json(output, translated)
    console.print(f"TRANSLATE {len(translated['segments'])} segments -> {output}")


@app.command()
def apply(
    input_path: Path = typer.Argument(..., exists=True, readable=True),
    translated_path: Path = typer.Argument(..., exists=True, readable=True),
    out: Path | None = typer.Option(None, "--out"),
    out_lang: str | None = typer.Option(None, "-out_lang", "--out-lang"),
) -> None:
    cfg = settings()
    target_lang = normalize_lang(out_lang or cfg.target_lang)
    output = out or default_output_path(input_path, target_lang)
    result = apply_translations(input_path, translated_path, output)
    status = "OK" if result["ok"] else f"WARN {result['issues']}"
    console.print(f"APPLY {status} -> {output}")


@app.command()
def run(
    input_path: Path = typer.Argument(..., exists=True, readable=True),
    in_lang: str | None = typer.Option(None, "-in_lang", "--in-lang"),
    out_lang: str | None = typer.Option(None, "-out_lang", "--out-lang"),
    lang: str | None = typer.Option(None, "-lang", "--lang"),
    output: Path | None = typer.Option(None, "--output"),
    no_move_done: bool = typer.Option(False, "--no-move-done"),
) -> None:
    cfg = settings()
    if in_lang:
        cfg.source_lang = normalize_lang(in_lang)
    if out_lang or lang:
        cfg.target_lang = normalize_lang(out_lang or lang or cfg.target_lang)
    files = discover_inputs(input_path)
    if not files:
        console.print("[yellow]처리할 .md 또는 .docx 파일이 없습니다.[/yellow]")
        return
    failures: list[tuple[Path, str]] = []
    for index, source in enumerate(files, start=1):
        console.print(f"[{index}/{len(files)}] RUN {source}")
        try:
            segments = extract_segments(source, cfg)
            work_dir = default_work_file(source, "segments.json", cfg).parent
            segments_path = work_dir / "segments.json"
            translated_path = work_dir / "translated.json"
            write_json(segments_path, segments)
            translated = translate_payload(segments, cfg)
            write_json(translated_path, translated)
            out_path = output if output and len(files) == 1 else default_output_path(source, cfg.target_lang, input_path if input_path.is_dir() else None)
            result = apply_translations(source, translated_path, out_path)
            if not result["ok"]:
                console.print(f"[yellow]VERIFY WARN {source}: {result['issues']}[/yellow]")
            if not no_move_done and source.parent.name.lower() == "input":
                moved = move_done(source)
                console.print(f"DONE 이동: {moved}")
            console.print(f"OUTPUT {out_path}")
        except Exception as exc:
            failures.append((source, str(exc)))
            console.print(f"[red]FAILED {source}: {exc}[/red]")
    if failures:
        console.print("[red]실패 목록[/red]")
        for source, message in failures:
            console.print(f"- {source}: {message}")
        raise typer.Exit(code=1)


@tm_app.command("import")
def tm_import(
    csv_path: Path = typer.Argument(..., exists=True, readable=True),
    in_lang: str | None = typer.Option(None, "-in_lang", "--in-lang"),
    out_lang: str | None = typer.Option(None, "-out_lang", "--out-lang"),
) -> None:
    cfg = settings()
    source_lang = normalize_lang(in_lang or cfg.source_lang)
    target_lang = normalize_lang(out_lang or cfg.target_lang)
    count = 0
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            src = (row.get("source") or row.get("src") or row.get("term") or "").strip()
            tgt = (row.get("target") or row.get("tgt") or row.get("translation") or "").strip()
            if src and tgt:
                tm_put(Path(cfg.tm_db_path), src, tgt, "import", source_lang, target_lang)
                count += 1
    console.print(f"TM import {count} rows")


if __name__ == "__main__":
    app()
